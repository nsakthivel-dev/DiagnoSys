import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc = docx.Document('RareDiseaseReport.docx')
with open('rare_disease_data_verbose.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + '\n')
    
    for table in doc.tables:
        for row in table.rows:
            f.write(' | '.join(cell.text for cell in row.cells) + '\n')
